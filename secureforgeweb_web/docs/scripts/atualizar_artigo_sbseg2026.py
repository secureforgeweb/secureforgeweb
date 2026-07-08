"""Insere capítulo Trabalhos Relacionados no artigo SBSEG2026 e renumera seções."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOC_PATH = Path(__file__).resolve().parent.parent / "SBSEG2026_SecureForge_Web.docx"

RENUMBER_MAP = {
    "2. Arquitetura da SecureForge Web": "3. Arquitetura da SecureForge Web",
    "2.1. Gestão de Aplicações e Análises": "3.1. Gestão de Aplicações e Análises",
    "2.2. Assessor HTTP (Headers de Segurança)": "3.2. Assessor HTTP (Headers de Segurança)",
    "2.3. Assessor Git (Análise Estática de Repositório)": "3.3. Assessor Git (Análise Estática de Repositório)",
    "2.4. Assistente IA por Usuário": "3.4. Assistente IA por Usuário",
    "2.5. Evidências e Rastreabilidade": "3.5. Evidências e Rastreabilidade",
    "2.6. Achados, Postura e Governança Administrativa": "3.6. Achados, Postura e Governança Administrativa",
    "3. Limitações": "4. Limitações",
    "4. Avaliação Experimental": "5. Avaliação Experimental",
    "4.1. Ambiente e Metodologia": "5.1. Ambiente e Metodologia",
    "4.2. Resultados": "5.2. Resultados",
    "5. Conclusão": "6. Conclusão",
    "6. Agradecimentos": "7. Agradecimentos",
}

RELATED_WORK_BLOCKS: list[tuple[str, str]] = [
    (
        "Heading 1",
        "2. Trabalhos Relacionados",
    ),
    (
        "Normal",
        "A literatura sobre segurança de aplicações web abrange desde padrões normativos "
        "e metodologias de teste até plataformas automatizadas de varredura e modelos "
        "quantitativos de postura. Esta seção posiciona a SecureForge Web em relação a "
        "esses trabalhos, destacando convergências e lacunas que motivam a proposta apresentada.",
    ),
    (
        "Heading 2",
        "2.1. Padrões OWASP e checklists de hardening",
    ),
    (
        "Normal",
        "O OWASP Application Security Verification Standard (ASVS) [2] estabelece "
        "requisitos técnicos verificáveis para aplicações web, enquanto o Web Security "
        "Testing Guide (WSTG) [3] orienta testes de segurança de forma metodológica. "
        "O Developer Guide da OWASP consolida checklists por domínio (autenticação, "
        "sessão, controle de acesso, validação de entrada, criptografia e configuração). "
        "Trabalhos como o de CIS [11] demonstram que diretrizes OWASP podem ser "
        "aplicadas ao longo do ciclo de desenvolvimento como checklist, com validação "
        "posterior por auditoria de código e testes de penetração. Esses artefatos "
        "definem o que verificar, mas não oferecem, por si só, um fluxo operacional "
        "integrado de cadastro de aplicações, wizard guiado, gestão de achados e "
        "relatório consolidado — lacuna que a SecureForge Web endereça.",
    ),
    (
        "Heading 2",
        "2.2. Avaliação quantitativa baseada em ASVS",
    ),
    (
        "Normal",
        "Wen e Katt [12] propõem modelo quantitativo de avaliação de segurança web "
        "fundamentado no ASVS, expressando o nível de segurança por métricas computacionais "
        "e análise de pontos fortes e fracos. Revniuk et al. [13] descrevem sistema de "
        "informação para avaliação quantitativa ASVS com pesos de especialistas, lógica fuzzy "
        "e visualização de resultados — arquiteturalmente próximo ao dashboard de postura "
        "da SecureForge Web. Em linha semelhante, estudos sobre métricas de security "
        "assurance [14] discutem como transformar conformidade qualitativa em indicadores "
        "mensuráveis. A diferença central da proposta aqui apresentada é combinar score de "
        "postura com automação assistida (HTTP, Git e IA) e validação humana explícita no "
        "wizard, em vez de depender exclusivamente de julgamento de especialistas ou de "
        "modelos puramente estáticos.",
    ),
    (
        "Heading 2",
        "2.3. Automação, DAST e ferramentas de hardening",
    ),
    (
        "Normal",
        "Ferramentas de Dynamic Application Security Testing (DAST), como o OWASP ZAP, "
        "automatizam a descoberta de vulnerabilidades em escala [15], porém tendem a "
        "produzir alto volume de alertas e exigem triagem especializada. Estudos em "
        "pipelines CI/CD [10] mapeiam verificações de segurança contínuas, enquanto "
        "trabalhos sobre operationalização do ASVS em CI/CD [16] buscam testes "
        "reproduzíveis em ambientes industriais. No domínio de headers HTTP, ferramentas "
        "especializadas (ex.: analisadores de HSTS, CSP e políticas de origem cruzada) "
        "oferecem verificação pontual e integração a pipelines. Projetos open source "
        "recentes — como frameworks de avaliação black-box com nota A–F e relatório PDF "
        "[17] ou plataformas de hardening com mapeamento OWASP Top 10 [18] — aproximam-se "
        "da automação parcial, mas não integram checklist guiado por item, evidências "
        "rastreáveis por resposta, assistente de IA configurável por usuário e benchmark "
        "administrativo entre operadores.",
    ),
    (
        "Heading 2",
        "2.4. Maturidade de processo e posicionamento da SecureForge Web",
    ),
    (
        "Normal",
        "Modelos de maturidade como o OWASP SAMM [19] e o BSIMM [20] avaliam a capacidade "
        "organizacional de segurança de software, não o hardening pontual de uma aplicação "
        "específica. A SecureForge Web ocupa nicho complementar: ferramenta operacional para "
        "equipes pequenas, laboratórios e projetos acadêmicos que necessitam de processo "
        "estruturado de revisão AppSec sem scanner enterprise. Em síntese, o trabalho se "
        "diferencia por unir (i) checklist OWASP guiado com 24 itens em nove categorias; "
        "(ii) três escopos de automação assistida com revisão humana obrigatória; "
        "(iii) persistência de evidências auditáveis; (iv) dashboard e exportação PDF; "
        "(v) configuração de IA por perfil e comparação administrativa entre análises — "
        "combinação não encontrada de forma integrada nos trabalhos relacionados revisados.",
    ),
]

NEW_REFERENCES = [
    (
        "CIS (2013). Developing a Secure Web Application Using OWASP Guidelines. "
        "Computer and Information Science, 2(4), 137–148. "
        "https://doi.org/10.5539/cis.v2n4p137. Acessado em: 07-07-2026."
    ),
    (
        "Wen, S.-F. and Katt, B. (2023). A quantitative security evaluation and analysis "
        "model for web applications based on OWASP application security verification standard. "
        "Computers & Security, 135, 103532. https://doi.org/10.1016/j.cose.2023.103532. "
        "Acessado em: 07-07-2026."
    ),
    (
        "Revniuk, O. et al. (2025). Development of an information system for the quantitative "
        "assessment of web application security based on the OWASP ASVS standard. "
        "Scientific Journal of the Ternopil Ivan Puluj National Technical University, 118(2), 56–65. "
        "https://sj.tntu.edu.ua/index.php/sjtntu/article/view/63. Acessado em: 07-07-2026."
    ),
    (
        "Wen, S.-F. et al. (2022). Developing Security Assurance Metrics to Support Quantitative "
        "Security Assurance Evaluation. Journal of Cybersecurity and Privacy, 2(3), 587–605. "
        "https://doi.org/10.3390/jcp2030030. Acessado em: 07-07-2026."
    ),
    (
        "Priyawati, D. et al. (2022). Website Vulnerability Testing and Analysis of Website "
        "Application Using OWASP. International Journal of Computer and Information System (IJCIS), "
        "3(3), 142–147. https://doi.org/10.29040/ijcis.v3i3.90. Acessado em: 07-07-2026."
    ),
    (
        "Operationalizing OWASP ASVS Level 1 in CI/CD (2024). Master's thesis, Tampere University. "
        "https://trepo.tuni.fi/handle/10024/233001. Acessado em: 07-07-2026."
    ),
    (
        "SHIELD Framework (2024). Dual-mode black-box security assessment framework. "
        "https://github.com/Georges034302/SHIELD-framework. Acessado em: 07-07-2026."
    ),
    (
        "StackSentry (2024). Automated web application security assessment framework. "
        "https://github.com/vickkykruz/sec_audit_framework. Acessado em: 07-07-2026."
    ),
    (
        "OWASP Foundation (2024). OWASP Software Assurance Maturity Model (SAMM). "
        "https://owasp.org/www-project-samm/. Acessado em: 07-07-2026."
    ),
    (
        "Synopsys (2024). Building Security In Maturity Model (BSIMM). "
        "https://www.synopsys.com/software-integrity/software-security-services/bsimm.html. "
        "Acessado em: 07-07-2026."
    ),
]


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def insert_blocks_after(paragraph: Paragraph, blocks: list[tuple[str, str]]) -> Paragraph:
    current = paragraph
    for style, text in blocks:
        current = insert_paragraph_after(current, text=text, style=style)
    return current


def main() -> None:
    doc = Document(str(DOC_PATH))

    anchor = None
    anchor_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "2. Arquitetura da SecureForge Web":
            anchor = p
            anchor_index = i
            break
    if anchor is None:
        raise RuntimeError("Âncora '2. Arquitetura' não encontrada.")

    insert_after = doc.paragraphs[anchor_index - 1] if anchor_index > 0 else doc.paragraphs[0]
    insert_blocks_after(insert_after, RELATED_WORK_BLOCKS)

    for p in doc.paragraphs:
        t = p.text.strip()
        if t in RENUMBER_MAP:
            p.text = RENUMBER_MAP[t]
        if "github.com/margefson/secureforgeweb" in t:
            p.text = t.replace(
                "https://github.com/margefson/secureforgeweb",
                "https://github.com/secureforgeweb/secureforgeweb",
            )

    ref_anchor = None
    ref_index = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Anais Estendidos do SBSeg"):
            ref_anchor = p
            ref_index = i
            break

    if ref_anchor is not None:
        insert_before = doc.paragraphs[ref_index - 1] if ref_index > 0 else ref_anchor
        current = insert_before
        for ref_text in NEW_REFERENCES:
            current = insert_paragraph_after(current, text=ref_text, style="Normal")
    else:
        last = doc.paragraphs[-1]
        current = last
        for ref_text in NEW_REFERENCES:
            current = insert_paragraph_after(current, text=ref_text, style="Normal")

    doc.save(str(DOC_PATH))
    print(f"Artigo atualizado: {DOC_PATH}")


if __name__ == "__main__":
    main()
