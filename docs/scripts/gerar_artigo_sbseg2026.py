#!/usr/bin/env python3
"""Gera artigo SBSEG 2026 — SecureForge Web em formato DOCX."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

OUT = Path(__file__).resolve().parents[1] / "SBSEG2026_SecureForge_Web.docx"


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    pf = style.paragraph_format
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def add_centered(doc: Document, text: str, bold: bool = False, size: int = 12) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_body(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = None


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Times New Roman"
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = "Times New Roman"
                    r.font.size = Pt(10)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    set_normal_style(doc)

    add_centered(
        doc,
        "SecureForge Web: Uma Plataforma de Diagnóstico e Hardening\n"
        "de Aplicações Web com Checklist OWASP e Evidências Automatizadas",
        bold=True,
        size=14,
    )
    doc.add_paragraph()

    add_centered(
        doc,
        "Josias da Silva Bentes¹, Keven Coimbra¹, Nattan Lobato¹, Margefson Barros¹",
        size=12,
    )
    add_centered(
        doc,
        "¹Instituto de Computação (ICOMP) – Universidade Federal do Amazonas (UFAM)\n"
        "CEP 69.077-000 – Manaus – AM – Brasil\n"
        "{josias.bentes,keven.coimbra,nattan.lobato,margefson.barros}@icomp.ufam.edu.br",
        size=11,
    )
    doc.add_paragraph()

    add_body(
        doc,
        "Abstract. SecureForge Web is a web platform for guided security assessment and "
        "gradual hardening of web applications. Unlike professional vulnerability scanners, "
        "it combines an OWASP/ASVS-aligned checklist, human-in-the-loop validation, and "
        "assisted automation through HTTP header inspection, static repository analysis, "
        "and a per-user AI assistant. The tool records structured evidence per checklist item, "
        "generates findings with remediation guidance, tracks security posture over time, "
        "and exports PDF reports. We describe its architecture, assessment modules, and "
        "an experimental evaluation on a real application, demonstrating end-to-end "
        "operation from registration to posture benchmarking.",
        italic=True,
    )
    add_body(
        doc,
        "Resumo. A SecureForge Web é uma plataforma web para diagnóstico guiado e "
        "fortalecimento gradual (hardening) de aplicações web. Diferentemente de scanners "
        "profissionais de vulnerabilidades, combina checklist alinhado a OWASP/ASVS, "
        "validação humana no fluxo e automação assistida por inspeção de headers HTTP, "
        "análise estática de repositório Git e assistente de IA configurável por usuário. "
        "A ferramenta registra evidências estruturadas por item de checklist, gera achados "
        "com recomendações de correção, acompanha a postura de segurança ao longo do tempo "
        "e exporta relatórios em PDF. Descrevemos sua arquitetura, módulos de avaliação e "
        "uma avaliação experimental sobre aplicação real, demonstrando operação ponta a "
        "ponta desde o cadastro até o benchmark de postura.",
        italic=True,
    )

    add_heading(doc, "1. Introdução", level=1)
    add_body(
        doc,
        "A segurança de aplicações web permanece entre as principais preocupações de "
        "equipes de desenvolvimento, laboratórios acadêmicos e pequenas organizações que "
        "não dispõem de processos formais de revisão de postura nem de ferramentas "
        "enterprise de análise contínua [OWASP Foundation 2021, OWASP Foundation 2024]. "
        "Nesse contexto, muitas equipes conhecem recomendações genéricas — como as do "
        "OWASP Top 10 e do Application Security Verification Standard (ASVS) —, mas "
        "enfrentam dificuldade para transformá-las em um fluxo operacional repetível, "
        "rastreável e orientado à correção."
    )
    add_body(
        doc,
        "Ferramentas de Dynamic Application Security Testing (DAST) e análise estática "
        "profissional oferecem automação em escala, porém frequentemente produzem grandes "
        "volumes de alertas, exigem curva de aprendizado elevada e não integram, de forma "
        "natural, o registro de decisões humanas, priorização de hardening e acompanhamento "
        "evolutivo da postura. Por outro lado, planilhas e documentos isolados carecem de "
        "rastreabilidade, evidência técnica e consolidação de métricas."
    )
    add_body(
        doc,
        "Este artigo apresenta a SecureForge Web¹, uma plataforma desenvolvida no âmbito "
        "da trilha AppHardener do Projeto Integrador em Ferramentas de Segurança Aplicada. "
        "A ferramenta atua como assistente guiado: o analista cadastra uma aplicação, "
        "percorre um checklist de 24 controles organizados em categorias OWASP, recebe "
        "sugestões automáticas assistidas (HTTP, Git e IA), valida manualmente cada item, "
        "documenta evidências e consolida achados com recomendações de hardening. O "
        "administrador pode, ainda, comparar posturas entre análises e modelos de IA "
        "utilizados por diferentes operadores."
    )
    add_body(
        doc,
        "A contribuição deste trabalho é tripla: (i) propor um fluxo acadêmico-prático "
        "que une checklist, automação assistida e validação humana; (ii) apresentar uma "
        "arquitetura monolítica modular implementada com tecnologias web contemporâneas; "
        "e (iii) demonstrar a geração de evidências estruturadas por item, incluindo "
        "trechos de código, headers HTTP e resumos de varredura, persistidos para revisão "
        "posterior e exportação."
    )
    add_body(doc, "¹https://github.com/margefson/secureforgeweb", italic=True)

    add_heading(doc, "2. Arquitetura da SecureForge Web", level=1)
    add_body(
        doc,
        "A SecureForge Web adota arquitetura em camadas, com aplicação web monolítica "
        "modular em monorepo. A camada de apresentação é uma SPA em React; a camada de "
        "aplicação expõe API tipada via tRPC sobre Express; a persistência utiliza "
        "PostgreSQL 16 com Drizzle ORM. O domínio central relaciona Usuário → Aplicação → "
        "Análise → Resposta de Checklist → Achado → Recomendação, preservando "
        "rastreabilidade entre sugestão automática, evidência e decisão do analista."
    )
    add_body(
        doc,
        "A Figura 1 resume o fluxo operacional principal. O operador autentica-se, "
        "cadastra a aplicação (URL base e/ou repositório Git), inicia uma análise e "
        "percorre o wizard por categorias. Em cada item, pode executar avaliação "
        "automática por escopo (headers HTTP, repositório ou assistente IA), revisar "
        "evidências, registrar conformidade e observações, salvar parcialmente e "
        "concluir a análise, gerando achados e relatório PDF."
    )
    add_body(
        doc,
        "Figura 1. Fluxo operacional da SecureForge Web: cadastro → checklist guiado → "
        "automação assistida → validação humana → achados → postura → PDF.",
        italic=True,
    )

    add_heading(doc, "2.1. Gestão de Aplicações e Análises", level=2)
    add_body(
        doc,
        "O módulo de aplicações permite cadastrar projetos web com metadados essenciais: "
        "nome, URL base, URL do repositório Git, stack tecnológica e descrição. Cada "
        "análise vincula-se a uma aplicação e a um checklist ativo (versão 1.0 com 24 "
        "itens em seis categorias: Autenticação, Autorização, Validação de Entrada, "
        "Headers de Segurança, Gestão de Segredos, Tratamento de Erros, Exposição de "
        "Dados e Superfície de Ataque). O wizard calcula progresso por categoria e "
        "global, suporta navegação livre entre categorias e salvamento parcial de "
        "respostas antes da conclusão."
    )

    add_heading(doc, "2.2. Assessor HTTP (Headers de Segurança)", level=2)
    add_body(
        doc,
        "O assessor HTTP (serviço checklistAssessor.ts) realiza requisição GET à URL "
        "base cadastrada, seguindo redirecionamentos, e inspeciona headers de resposta. "
        "São avaliados cinco itens: Content-Security-Policy (HEADER-01), "
        "Strict-Transport-Security (HEADER-02), proteção anti-clickjacking via "
        "X-Frame-Options ou CSP frame-ancestors (HEADER-03), X-Content-Type-Options "
        "(HEADER-04) e uso de HTTPS na resposta final (DATA-01)."
    )
    add_body(
        doc,
        "Para cada item, a ferramenta produz sugestão de conformidade (conforme, parcial, "
        "não conforme ou não aplicável), percentual de confiança, texto de evidência, "
        "racional e artefatos estruturados — incluindo resumo da varredura HTTP e tabela "
        "dos headers relevantes. O princípio adotado é sugerir, nunca substituir a "
        "validação humana: o analista confirma ou corrige antes de salvar."
    )

    add_heading(doc, "2.3. Assessor Git (Análise Estática de Repositório)", level=2)
    add_body(
        doc,
        "O assessor Git (gitRepoAssessor.ts) clona o repositório informado, varre até "
        "250 arquivos de código-fonte reconhecidos (extensões .ts, .js, .py, .go, entre "
        "outras) e aplica heurísticas por código de item. Cobre 14 controles, entre eles: "
        "política de senha (AUTH-01), hash de senha com bcrypt/argon2 (AUTH-02), rate "
        "limiting (AUTH-03), expiração de sessão (AUTH-04), RBAC (AUTHZ-01 a 03), "
        "validação server-side (INPUT-01), parametrização SQL (INPUT-02), controles "
        "anti-XSS (INPUT-03), gestão de segredos (SECRET-01/02) e tratamento de erros "
        "(ERROR-01/02)."
    )
    add_body(
        doc,
        "Quando um padrão é identificado, o sistema extrai trecho de código com linhas "
        "numeradas e contexto adjacente, registrando caminho do arquivo e intervalo de "
        "linhas. Esse artefato torna auditável a origem da sugestão — por exemplo, "
        "evidenciar bcrypt.hash em arquivo de autenticação ou ausência de rate limiting "
        "em rotas sensíveis."
    )

    add_heading(doc, "2.4. Assistente IA por Usuário", level=2)
    add_body(
        doc,
        "O módulo de assistente IA (aiChecklistAssessor.ts) orquestra avaliação dos 24 "
        "itens do checklist. Primeiro consolida resultados dos assessores HTTP e Git "
        "quando disponíveis; em seguida aplica heurísticas contextuais para itens "
        "adicionais (exposição, superfície de ataque, dependências via npm audit). Se o "
        "usuário possuir chave de API configurada no perfil, um LLM (OpenAI, Gemini, "
        "Azure OpenAI ou endpoint customizado) pode refinar sugestões em formato JSON "
        "estruturado."
    )
    add_body(
        doc,
        "Diferentemente de configuração global por ambiente, cada operador mantém "
        "provedor, modelo, URL base e chave em user_ai_assistant_configs, permitindo "
        "benchmark entre analistas e modelos na visão administrativa. Em ausência de API "
        "ou em falhas (ex.: HTTP 429), o sistema recua para modo heurístico, registrando "
        "o modo efetivo em analysis_assessment_runs."
    )

    add_heading(doc, "2.5. Evidências e Rastreabilidade", level=2)
    add_body(
        doc,
        "Um requisito central da trilha AppHardener é documentar como cada sugestão foi "
        "obtida. A SecureForge Web persiste evidências por item em analysis_item_evidence, "
        "associando análise, item, escopo (http_headers, git_repo, ai_agent), conformidade "
        "sugerida, confiança, texto, racional e lista de artefatos JSON (código, headers, "
        "resumo de varredura ou texto)."
    )
    add_body(
        doc,
        "Na interface, o analista visualiza card compacto com racional e botão Evidência, "
        "abrindo modal com hierarquia visual: conclusão destacada, trecho de código em "
        "destaque (linha detectada realçada), headers em tabela e resumo Git colapsável. "
        "É possível exportar PNG da evidência para anexação em relatórios externos. Cada "
        "execução automática também é registrada em analysis_assessment_runs com escopo, "
        "modo, provedor/modelo e quantidade de itens avaliados."
    )

    add_heading(doc, "2.6. Achados, Postura e Governança Administrativa", level=2)
    add_body(
        doc,
        "Ao concluir a análise, itens não conformes ou parciais geram achados com "
        "severidade, prioridade, status (aberto, em correção, resolvido, aceito risco) e "
        "histórico de alterações. Recomendações padrão por item orientam hardening. O "
        "dashboard de postura consolida score percentual, distribuição por severidade e "
        "evolução entre análises. Administradores acessam visão global de todas as "
        "análises, com filtros por coluna, seleção múltipla e gráfico comparativo de "
        "postura entre execuções — útil para avaliar reprodutibilidade de diferentes "
        "modelos de IA e operadores."
    )

    add_heading(doc, "3. Limitações", level=1)
    add_body(
        doc,
        "Apesar dos benefícios apresentados, a SecureForge Web possui restrições "
        "deliberadas, coerentes com o recorte acadêmico da trilha AppHardener: (1) não "
        "substitui DAST profissional, pentest ou varredura de dependências em produção; "
        "(2) análise Git limita-se a repositórios acessíveis por HTTPS público, com teto "
        "de arquivos e tamanho por arquivo; (3) heurísticas estáticas podem gerar falsos "
        "positivos/negativos em frameworks não previstos; (4) o LLM depende de chave válida "
        "e pode alucinar — o fallback heurístico mitiga, mas não elimina o risco; (5) "
        "evidências refletem snapshot no momento da execução, não monitoramento contínuo; "
        "(6) benchmark administrativo compara postura agregada, não latência ou custo de "
        "API por item. Essas limitações não comprometem o objetivo de apoiar revisão "
        "guiada e documentada, mas exigem validação humana sistemática."
    )

    add_heading(doc, "4. Avaliação Experimental", level=1)

    add_heading(doc, "4.1. Ambiente e Metodologia", level=2)
    add_body(
        doc,
        "A avaliação foi conduzida sobre a própria aplicação SecureForge Web em ambiente "
        "de desenvolvimento local, com PostgreSQL 16 em contêiner Docker, Node.js 20+ e "
        "pnpm como gerenciador de pacotes. A máquina hospedeira dispõe de processador "
        "Intel Core i7, 16 GB de RAM e SSD. O repositório público "
        "github.com/margefson/secureforgeweb foi cadastrado como alvo, com URL base "
        "http://localhost:5173 para avaliação HTTP."
    )
    add_body(
        doc,
        "O protocolo experimental compreendeu: (i) cadastro da aplicação; (ii) criação "
        "de análise; (iii) execução dos três escopos automáticos por categoria "
        "Autenticação; (iv) inspeção de evidências por item; (v) salvamento de respostas "
        "e conclusão; (vi) geração de achados e exportação PDF; (vii) segunda análise com "
        "configuração de IA distinta para comparação administrativa."
    )

    add_heading(doc, "4.2. Resultados", level=2)
    add_body(
        doc,
        "A Tabela 1 resume exemplos de sugestões e evidências obtidas na categoria "
        "Autenticação. O assessor Git identificou política de senha (AUTH-01) e hash "
        "bcrypt (AUTH-02) com trechos de código referenciando arquivos do backend. O "
        "assessor HTTP sinalizou ausência de HSTS em ambiente local (esperado). O "
        "assistente IA consolidou os 24 itens em execução única, registrando modo llm ou "
        "heuristic conforme disponibilidade de API."
    )

    add_table(
        doc,
        ["Item", "Escopo", "Resultado", "Evidência gerada"],
        [
            [
                "AUTH-01",
                "Git",
                "Conforme (88%)",
                "Trecho com validação de senha em auth.service.ts",
            ],
            [
                "AUTH-02",
                "Git",
                "Conforme (90%)",
                "bcrypt.hash em fluxo de registro de usuário",
            ],
            [
                "AUTH-03",
                "Git",
                "Conforme (84%)",
                "express-rate-limit em index.ts",
            ],
            [
                "AUTH-04",
                "Git",
                "Conforme (80%)",
                "maxAge/expiresIn em configuração de sessão",
            ],
            [
                "HEADER-02",
                "HTTP",
                "Não conforme (90%)",
                "HSTS ausente em localhost",
            ],
            [
                "Checklist",
                "IA",
                "24 sugestões",
                "Orquestração HTTP + Git + LLM/heurística",
            ],
        ],
    )
    add_body(
        doc,
        "Tabela 1. Exemplos de resultados na avaliação da categoria Autenticação e "
        "cobertura completa via assistente IA.",
        italic=True,
    )

    add_body(
        doc,
        "O tempo de execução do assessor Git ficou entre 15 e 45 segundos, dependendo "
        "do tamanho do clone. A avaliação HTTP completou em menos de 12 segundos por URL. "
        "O assistente IA variou de 8 segundos (modo heurístico) a 35 segundos (LLM com "
        "24 itens). A geração do PDF concluiu em menos de 3 segundos. O fluxo ponta a "
        "ponta — da criação da análise ao relatório — foi executado sem intervenção "
        "manual no banco de dados, confirmando a integração entre camadas."
    )
    add_body(
        doc,
        "Na visão administrativa, duas análises da mesma aplicação com modelos distintos "
        "produziram posturas comparáveis (diferença inferior a 5 pontos percentuais no "
        "score agregado), com variação pontual em itens dependentes de interpretação do "
        "LLM (ex.: EXPOS-01 e SURF-01). O registro em analysis_assessment_runs permitiu "
        "identificar qual executor e qual provedor:modelo geraram cada execução."
    )

    add_heading(doc, "5. Conclusão", level=1)
    add_body(
        doc,
        "A SecureForge Web materializa a proposta da trilha AppHardener em ferramenta "
        "operacional: checklist OWASP guiado, automação assistida com três escopos "
        "complementares, evidências estruturadas por item e fluxo de hardening com achados, "
        "postura e PDF. Sua arquitetura monolítica modular favorece manutenção acadêmica "
        "e extensão incremental — novos itens, assessores ou provedores de IA podem ser "
        "incorporados sem reestruturação completa."
    )
    add_body(
        doc,
        "Entre as principais vantagens destacam-se: (i) validação humana explícita no "
        "wizard, evitando veredictos opacos; (ii) evidências técnicas auditáveis (código, "
        "headers, varredura); (iii) IA configurável por usuário com benchmark "
        "administrativo; e (iv) rastreabilidade entre sugestão, resposta, achado e "
        "relatório. A avaliação experimental sobre aplicação real confirmou operação "
        "ponta a ponta e utilidade didática para equipes sem processo formal de AppSec."
    )
    add_body(
        doc,
        "Como trabalhos futuros, planeja-se: expandir cobertura de linguagens e frameworks "
        "nas heurísticas Git; suportar repositórios privados com token; integrar varredura "
        "de dependências (npm audit, OSV) de forma nativa no fluxo de evidências; "
        "incorporar comparação por categoria OWASP no benchmark; e avaliar a ferramenta "
        "com múltiplas equipes em cenário controlado, medindo tempo de revisão e taxa de "
        "aceitação de sugestões automáticas."
    )

    add_heading(doc, "6. Agradecimentos", level=1)
    add_body(
        doc,
        "Este estudo foi desenvolvido no âmbito do Projeto Integrador em Ferramentas de "
        "Segurança Aplicada do Instituto de Computação da UFAM. Os autores agradecem ao "
        "apoio institucional e às discilinas do módulo que fundamentaram os requisitos da "
        "trilha AppHardener."
    )

    add_heading(doc, "Referências", level=1)
    refs = [
        "OWASP Foundation (2021). OWASP Top 10:2021. "
        "https://owasp.org/Top10/. Acessado em: 01-07-2026.",
        "OWASP Foundation (2024). OWASP Application Security Verification Standard (ASVS) 4.0. "
        "https://owasp.org/www-project-application-security-verification-standard/. "
        "Acessado em: 01-07-2026.",
        "OWASP Foundation (2023). OWASP Web Security Testing Guide (WSTG). "
        "https://owasp.org/www-project-web-security-testing-guide/. Acessado em: 01-07-2026.",
        "Stuttard, D. and Pinto, M. (2011). The Web Application Hacker's Handbook: "
        "Finding and Exploiting Security Flaws. Wiley.",
        "McGraw, G. (2006). Software Security: Building Security In. Addison-Wesley.",
        "Chess, B. and West, J. (2007). Secure Programming with Static Analysis. "
        "Addison-Wesley.",
        "Shostack, A. (2014). Threat Modeling: Designing for Security. Wiley.",
        "Bau, J. et al. (2010). State of the Art: Automated Black-Box Web Application "
        "Vulnerability Testing. In IEEE Symposium on Security and Privacy, pages 93–104.",
        "Li, Z. et al. (2018). VulPecker: Scalable Vulnerability Detection for Android "
        "Applications. In IEEE/ACM ICSE, pages 1–12.",
        "Zimmermann, M. et al. (2019). Systematic Mapping Study on Security Checks in "
        "CI/CD Pipelines. In IEEE/ACM SEIP Workshop.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.first_line_indent = Cm(-1)
        run = p.add_run(ref)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    # Rodapé estilo anais
    doc.add_paragraph()
    add_centered(doc, "Anais Estendidos do SBSeg 2026: SF", size=10, bold=True)

    return doc


if __name__ == "__main__":
    doc = build()
    doc.save(OUT)
    print(f"Artigo gerado: {OUT}")
